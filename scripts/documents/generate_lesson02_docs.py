from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "public" / "resources" / "lesson-02"
OUT.mkdir(parents=True, exist_ok=True)

# compact_reference_guide preset with named school-material overrides:
# - A4 instead of US Letter; portrait worksheets and landscape teacher sheets.
# - S-Core Dream for every Korean run.
# - 10 pt minimum student activity text and generous touch/marking areas.
# - Warm paper, forest green and restrained brick accents for a Korean school handout.
FONT = "S-Core Dream 4 Regular"
FONT_KO = "에스코어 드림 4 Regular"
COLORS = {
    "ink": "24322F",
    "muted": "66736E",
    "green": "315B50",
    "green_dark": "23473F",
    "sage": "E7EFEA",
    "sage_light": "F3F7F4",
    "paper": "FCFAF4",
    "sand": "F3EBDD",
    "brick": "A45643",
    "brick_light": "F4E7E1",
    "line": "AAB9B2",
    "white": "FFFFFF",
}

DXA_PER_MM = 1440 / 25.4
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 90, "bottom": 90, "start": 120, "end": 120}


STORIES = [
    {
        "title": "해와 달이 된 오누이",
        "theme": "용기와 지혜",
        "characters": [
            ("오누이", True), ("어머니", True), ("호랑이", True),
            ("제비", False), ("임금", False),
        ],
        "backgrounds": [
            ("산길", True), ("오누이의 집", True), ("큰 나무", True),
            ("하늘", True), ("궁궐", False), ("용궁", False),
        ],
        "events": [
            "어머니가 장에 떡을 팔러 떠나요.",
            "산길에서 호랑이가 어머니를 만나요.",
            "호랑이가 오누이의 집으로 찾아와요.",
            "오누이가 호랑이를 피해 큰 나무로 올라가요.",
            "오누이가 하늘을 향해 도와 달라고 빌어요.",
            "오누이는 해와 달이 되어 하늘을 밝혀요.",
        ],
        "plans": [
            {
                "who": ["어머니", "오누이"],
                "where": ["집 앞", "마을길"],
                "feeling": ["활기차게", "조금 걱정되게"],
                "look": ["떡 바구니가 보이게", "손 흔드는 모습"],
            },
            {
                "who": ["어머니", "호랑이"],
                "where": ["산길", "고갯길"],
                "feeling": ["무섭게", "긴장되게"],
                "look": ["호랑이를 크게", "놀란 얼굴이 보이게"],
            },
            {
                "who": ["오누이", "호랑이"],
                "where": ["집 안", "대문 앞"],
                "feeling": ["어둡게", "조마조마하게"],
                "look": ["문 앞 호랑이", "귀 기울이는 오누이"],
            },
            {
                "who": ["오누이", "호랑이"],
                "where": ["큰 나무", "집 마당"],
                "feeling": ["다급하게", "용감하게"],
                "look": ["오르는 모습", "나무 아래 호랑이"],
            },
            {
                "who": ["오누이"],
                "where": ["나무 위", "하늘 아래"],
                "feeling": ["간절하게", "희망차게"],
                "look": ["두 손을 모으게", "하늘을 바라보게"],
            },
            {
                "who": ["오누이", "해와 달"],
                "where": ["하늘", "구름 사이"],
                "feeling": ["밝게", "따뜻하게", "신비롭게"],
                "look": ["해와 달을 크게", "빛이 퍼지게"],
            },
        ],
    },
    {
        "title": "흥부와 놀부",
        "theme": "나눔과 배려",
        "characters": [
            ("흥부", True), ("놀부", True), ("제비", True),
            ("흥부의 가족", True), ("호랑이", False), ("선녀", False),
        ],
        "backgrounds": [
            ("흥부네 집", True), ("놀부네 집", True), ("제비집", True),
            ("박이 열린 마당", True), ("용궁", False), ("궁궐", False),
        ],
        "events": [
            "흥부와 놀부는 서로 다른 마음으로 살아가요.",
            "흥부가 다친 제비의 다리를 정성껏 고쳐 줘요.",
            "돌아온 제비가 흥부에게 박씨를 물어다 줘요.",
            "흥부 가족은 박을 타고 큰 도움을 받아요.",
            "놀부도 흥부를 따라 하며 박씨를 얻으려 해요.",
            "놀부는 잘못을 깨닫고 형제는 사이좋게 지내요.",
        ],
        "plans": [
            {
                "who": ["흥부", "놀부", "두 가족"],
                "where": ["두 집 앞", "마을길"],
                "feeling": ["따뜻하게", "서로 다르게"],
                "look": ["두 집이 보이게", "표정이 다르게"],
            },
            {
                "who": ["흥부", "다친 제비"],
                "where": ["제비집 아래", "흥부네 마당"],
                "feeling": ["걱정되게", "다정하게"],
                "look": ["다친 다리가 보이게", "돌보는 손을 크게"],
            },
            {
                "who": ["흥부", "돌아온 제비"],
                "where": ["흥부네 마당", "제비집 앞"],
                "feeling": ["반갑게", "신나게"],
                "look": ["부리의 박씨", "받는 손을 크게"],
            },
            {
                "who": ["흥부 가족"],
                "where": ["박이 열린 마당", "흥부네 집"],
                "feeling": ["놀랍게", "기쁘게"],
                "look": ["갈라지는 박", "둘러선 가족"],
            },
            {
                "who": ["놀부", "제비"],
                "where": ["놀부네 마당", "제비집 앞"],
                "feeling": ["욕심나게", "초조하게"],
                "look": ["박씨를 바라보게", "찡그린 표정"],
            },
            {
                "who": ["흥부", "놀부", "두 가족"],
                "where": ["집 앞마당", "마을"],
                "feeling": ["따뜻하게", "즐겁게"],
                "look": ["마주 잡은 손", "나누는 음식"],
            },
        ],
    },
    {
        "title": "의좋은 형제",
        "theme": "서로 아끼는 마음",
        "characters": [
            ("형", True), ("동생", True), ("임금", False),
            ("호랑이", False), ("제비", False),
        ],
        "backgrounds": [
            ("논", True), ("형의 집", True), ("동생의 집", True),
            ("밤길", True), ("궁궐", False), ("산속 동굴", False),
        ],
        "events": [
            "형과 동생은 각자 농사를 지으며 사이좋게 살아요.",
            "가을이 되어 볏단을 똑같이 나누어요.",
            "형은 동생에게 더 주려고 밤에 볏단을 옮겨요.",
            "동생도 형에게 더 주려고 밤에 볏단을 옮겨요.",
            "다음 날에도 볏단 수가 그대로여서 두 사람은 궁금해해요.",
            "밤길에서 만난 형제는 서로의 마음을 알고 꼭 안아요.",
        ],
        "plans": [
            {
                "who": ["형", "동생"],
                "where": ["논", "집 가까운 들"],
                "feeling": ["평화롭게", "부지런하게"],
                "look": ["함께 일하게", "논이 넓게 보이게"],
            },
            {
                "who": ["형", "동생"],
                "where": ["논", "볏단 앞"],
                "feeling": ["사이좋게", "뿌듯하게"],
                "look": ["볏단을 나누게", "볏단이 많이 보이게"],
            },
            {
                "who": ["형", "잠든 동생"],
                "where": ["밤길", "동생의 집 앞"],
                "feeling": ["조용하게", "따뜻하게"],
                "look": ["볏단을 지게", "달빛 길이 보이게"],
            },
            {
                "who": ["동생", "잠든 형"],
                "where": ["밤길", "형의 집 앞"],
                "feeling": ["조용하게", "따뜻하게"],
                "look": ["볏단을 지게", "발걸음이 보이게"],
            },
            {
                "who": ["형", "동생"],
                "where": ["각자의 집", "볏단 앞"],
                "feeling": ["궁금하게", "놀랍게"],
                "look": ["볏단을 세게", "고개를 갸웃하게"],
            },
            {
                "who": ["형", "동생"],
                "where": ["밤길", "논 가까운 길"],
                "feeling": ["감동적으로", "기쁘게"],
                "look": ["서로 꼭 안게", "볏단과 달이 보이게"],
            },
        ],
    },
]

# Each page shows A-F in this deliberately shuffled order.
# A=4th, B=1st, C=6th, D=3rd, E=2nd, F=5th; key is B-E-D-A-F-C.
DISPLAY_ORDER = [3, 0, 5, 2, 1, 4]
HANGUL_WORD_RE = re.compile(r"[가-힣]{2,}")


def mm_to_dxa(value: float) -> int:
    return round(value * DXA_PER_MM)


def keep_hangul_words(text: str) -> str:
    """Prevent renderers from breaking Korean words between syllables."""
    return HANGUL_WORD_RE.sub(lambda match: "\u2060".join(match.group(0)), text)


def set_run(run, *, size: float = 10.5, bold: bool = False, color: str = "ink"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key, name in (
        ("ascii", FONT), ("hAnsi", FONT), ("eastAsia", FONT_KO), ("cs", FONT)
    ):
        r_fonts.set(qn(f"w:{key}"), name)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(COLORS.get(color, color))


def style_paragraph(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 0,
    line: float = 1.25,
    keep_with_next: bool = False,
    keep_together: bool = True,
):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together
    fmt.widow_control = True
    return paragraph


def add_text(
    paragraph,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    color: str = "ink",
):
    run = paragraph.add_run(keep_hangul_words(text))
    set_run(run, size=size, bold=bold, color=color)
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), COLORS.get(fill, fill))


def set_cell_margins(cell, *, top=90, bottom=90, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = "line", size: int = 5, inside: bool = True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges.extend(["insideH", "insideV"])
    for edge in edges:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), COLORS.get(color, color))


def set_table_geometry(table, widths_mm: list[float] | tuple[float, ...], indent_dxa: int = TABLE_INDENT_DXA):
    """Set tblW/tblInd/tblGrid/tcW to matching fixed DXA values, including merged cells."""
    widths = [mm_to_dxa(value) for value in widths_mm]
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for tr in table._tbl.tr_lst:
        col_index = 0
        for tc in tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            span_node = tc_pr.find(qn("w:gridSpan"))
            span = int(span_node.get(qn("w:val"))) if span_node is not None else 1
            width = sum(widths[col_index: col_index + span])
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            col_index += span


def prevent_row_split(row, *, min_height_mm: float | None = None):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if min_height_mm is not None:
        row.height = Mm(min_height_mm)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def fill_cell(
    cell,
    text: str = "",
    *,
    size: float = 10.5,
    bold: bool = False,
    color: str = "ink",
    fill: str | None = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line: float = 1.2,
    margins: dict[str, int] | None = None,
):
    cell.text = ""
    p = style_paragraph(cell.paragraphs[0], align=align, line=line)
    if text:
        add_text(p, text, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell, **(margins or CELL_MARGIN))
    return cell


def setup_doc(title: str, *, landscape: bool = False, student: bool = False):
    doc = Document()
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Mm(297), Mm(210)
        section.top_margin = Mm(6)
        section.bottom_margin = Mm(6)
        section.left_margin = Mm(8)
        section.right_margin = Mm(8)
    else:
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = Mm(8)
        section.bottom_margin = Mm(8)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(4)

    doc.core_properties.author = "더몬스터학원 EDU STORY"
    doc.core_properties.title = title
    doc.core_properties.subject = "AI와 함께 만드는 우리 옛이야기 그림책 2차시"

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_KO)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 20, "ink", 0, 4),
        ("Heading 1", 16, "green", 18, 10),
        ("Heading 2", 13, "green", 14, 7),
        ("Heading 3", 12, "green_dark", 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), FONT)
        r_fonts.set(qn("w:hAnsi"), FONT)
        r_fonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(size)
        style.font.bold = name != "Title"
        style.font.color.rgb = RGBColor.from_string(COLORS[color])
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer
    p = style_paragraph(footer.paragraphs[0], align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    add_text(
        p,
        "AI와 함께 만드는 우리 옛이야기 그림책  |  더몬스터학원 EDU STORY",
        size=7.2 if landscape else (10.0 if student else 7.2),
        color="muted",
    )
    return doc


def add_masthead(
    doc,
    *,
    kicker: str,
    title: str,
    subtitle: str,
    marker: str,
    wide: bool = False,
    compact: bool = False,
):
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    pad = 70 if compact else 120
    fill_cell(left, fill="paper", margins={"top": pad, "bottom": pad, "start": 150, "end": 150})
    p = style_paragraph(left.paragraphs[0], after=1, keep_with_next=True)
    add_text(p, kicker, size=7.5 if compact else 10.0, bold=True, color="green")
    p = style_paragraph(left.add_paragraph(), after=1, line=1.1, keep_with_next=True)
    add_text(p, title, size=16.8 if compact else 19, bold=True, color="ink")
    p = style_paragraph(left.add_paragraph(), line=1.15)
    add_text(p, subtitle, size=8.3 if compact else 10.2, bold=True, color="brick")

    fill_cell(right, marker, size=9.0 if compact else 10.5, bold=True, color="green_dark", fill="sage", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [205, 76] if wide else [130, 57.5])
    set_table_borders(table, color="green", size=7, inside=True)
    prevent_row_split(table.rows[0], min_height_mm=17 if compact else 23)
    return table


def add_section_heading(doc, title: str, instruction: str, *, compact: bool = False):
    p = style_paragraph(
        doc.add_paragraph(),
        before=2 if compact else 4,
        after=1 if compact else 2,
        line=1.05 if compact else 1.15,
        keep_with_next=True,
    )
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), COLORS["brick"])
    borders.append(left)
    p_pr.append(borders)
    add_text(p, title, size=10.2 if compact else 11.8, bold=True, color="green_dark")
    add_text(p, f"  {instruction}", size=8.5 if compact else 10.2, color="muted")
    return p


def add_student_fields(doc, *, wide: bool = False):
    table = doc.add_table(rows=1, cols=3)
    labels = ["학년·반  __________", "모둠  ______", "이름  __________"]
    for index, label in enumerate(labels):
        fill_cell(
            table.cell(0, index), label,
            size=8.8 if wide else 10.0,
            bold=True,
            fill="sage_light",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            margins={"top": 35, "bottom": 35, "start": 70, "end": 70} if wide else None,
        )
    set_table_geometry(table, [93, 80, 105.5] if wide else [62.5, 50, 75])
    set_table_borders(table, color="line", size=4, inside=True)
    prevent_row_split(table.rows[0], min_height_mm=8 if wide else 10)


def add_student_sheet_header(doc, story, story_index: int):
    """Compact one-row masthead so one story version stays on one A4 sheet."""
    table = doc.add_table(rows=1, cols=4)
    title_cell = table.cell(0, 0)
    fill_cell(
        title_cell,
        fill="paper",
        margins={"top": 55, "bottom": 55, "start": 120, "end": 100},
    )
    p = style_paragraph(title_cell.paragraphs[0], after=0.5, line=1.0)
    add_text(p, "2차시  ", size=8.5, bold=True, color="brick")
    add_text(p, "우리 그림책 여섯 장면 기획하기", size=14.0, bold=True, color="ink")
    p = style_paragraph(title_cell.add_paragraph(), after=0, line=1.0)
    add_text(p, f"{story['title']}  ·  {story['theme']}  ·  이야기 {story_index} / 3", size=8.2, bold=True, color="green")

    fields = ["학년·반\n________", "모둠\n______", "이름\n________"]
    for col_index, label in enumerate(fields, start=1):
        fill_cell(
            table.cell(0, col_index),
            label,
            size=8.2,
            bold=True,
            color="green_dark",
            fill="sage",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.05,
            margins={"top": 40, "bottom": 40, "start": 45, "end": 45},
        )
    set_table_geometry(table, [157, 40, 35, 46.5])
    set_table_borders(table, color="green", size=6, inside=True)
    prevent_row_split(table.rows[0], min_height_mm=17)


def option_text(items: list[tuple[str, bool]]) -> str:
    return "   ".join(f"○ {label}" for label, _ in items)


def choice_line(cell, label: str, values: list[str]):
    p = style_paragraph(cell.add_paragraph(), after=0.3, line=1.0, keep_together=True)
    add_text(p, f"{label}  ", size=8.0, bold=True, color="green")
    add_text(p, "  ".join(f"○ {value}" for value in values), size=8.0, color="ink")


def add_scene_planning_grid(doc, story):
    """Six compact student-authored picture-book planning panels."""
    grid = doc.add_table(rows=2, cols=3)
    for scene_index, (event, plan) in enumerate(zip(story["events"], story["plans"])):
        row_index, col_index = divmod(scene_index, 3)
        cell = grid.cell(row_index, col_index)
        fill_cell(
            cell,
            fill="paper" if scene_index % 2 == 0 else "white",
            margins={"top": 40, "bottom": 40, "start": 70, "end": 70},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        p = style_paragraph(cell.paragraphs[0], after=0.5, line=1.0, keep_with_next=True)
        add_text(p, f"{scene_index + 1:02d} 장면  ", size=7.7, bold=True, color="brick")
        add_text(p, event, size=8.3, bold=True, color="ink")
        choice_line(cell, "누가?", plan["who"])
        choice_line(cell, "어디?", plan["where"])
        choice_line(cell, "느낌?", plan["feeling"])
        choice_line(cell, "모습?", plan["look"])

    set_table_geometry(grid, [92.83, 92.83, 92.84])
    set_table_borders(grid, color="green", size=5, inside=True)
    for row in grid.rows:
        prevent_row_split(row, min_height_mm=27)


def add_chosen_scene_sketch(doc):
    add_section_heading(
        doc,
        "4  한 장면을 그림으로 메모하기",
        "마음에 드는 장면 하나만 골라 선이나 모양으로 그려요.",
        compact=True,
    )
    panel = doc.add_table(rows=1, cols=2)
    choices, sketch = panel.rows[0].cells
    fill_cell(
        choices,
        fill="sage_light",
        margins={"top": 55, "bottom": 55, "start": 100, "end": 100},
    )
    choices.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = choices.paragraphs[0]
    add_text(p, "내가 고른 장면  ", size=8.3, bold=True, color="green_dark")
    add_text(p, "______번", size=8.3, bold=True, color="ink")
    p = style_paragraph(choices.add_paragraph(), before=2, after=0, line=1.1)
    add_text(p, "가장 크게 보여 줄 것  ", size=8.3, bold=True, color="green_dark")
    add_text(p, "○ 얼굴  ○ 몸짓  ○ 중요한 물건  ○ 넓은 장소", size=8.3, color="ink")
    p = style_paragraph(choices.add_paragraph(), before=2, after=0, line=1.1)
    add_text(p, "친구에게 장면을 말했어요  ○", size=8.3, bold=True, color="brick")

    fill_cell(
        sketch,
        "그림 메모  ·  잘 그리지 않아도 괜찮아요.",
        size=8.3,
        bold=True,
        color="muted",
        fill="paper",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line=1.1,
        margins={"top": 55, "bottom": 55, "start": 80, "end": 80},
    )
    set_table_geometry(panel, [94, 184.5])
    set_table_borders(panel, color="line", size=5, inside=True)
    prevent_row_split(panel.rows[0], min_height_mm=24)


def make_worksheet_pack() -> Path:
    doc = setup_doc("2차시 학생 활동지 3종", landscape=True, student=True)

    for page_index, story in enumerate(STORIES, start=1):
        if page_index > 1:
            doc.add_page_break()
        add_student_sheet_header(doc, story, page_index)

        add_section_heading(doc, "1  이야기 뼈대 확인", "나오는 인물과 장소에 ○표해요.", compact=True)
        options = doc.add_table(rows=1, cols=4)
        fill_cell(options.cell(0, 0), "인물", size=8.1, bold=True, color="green_dark", fill="sage", align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(options.cell(0, 1), option_text(story["characters"]), size=8.2, bold=True, line=1.08, margins={"top": 35, "bottom": 35, "start": 70, "end": 70})
        fill_cell(options.cell(0, 2), "장소", size=8.1, bold=True, color="green_dark", fill="sage", align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(options.cell(0, 3), option_text(story["backgrounds"]), size=8.2, bold=True, line=1.08, margins={"top": 35, "bottom": 35, "start": 70, "end": 70})
        set_table_geometry(options, [16, 121, 16, 125.5])
        set_table_borders(options, color="line", size=5, inside=True)
        prevent_row_split(options.rows[0], min_height_mm=10)

        add_section_heading(doc, "2  사건 순서 맞추기", "A-F를 먼저 일어난 일부터 배열해요.", compact=True)
        event_table = doc.add_table(rows=2, cols=6)
        for position, event_index in enumerate(DISPLAY_ORDER):
            row_index, side = divmod(position, 3)
            label_col = side * 2
            text_col = label_col + 1
            label = chr(65 + position)
            fill_cell(
                event_table.cell(row_index, label_col), label,
                size=9.0, bold=True, color="white", fill="green", align=WD_ALIGN_PARAGRAPH.CENTER,
                margins={"top": 35, "bottom": 35, "start": 35, "end": 35},
            )
            fill_cell(
                event_table.cell(row_index, text_col), story["events"][event_index],
                size=8.0, bold=True, fill="paper", line=1.08,
                margins={"top": 45, "bottom": 45, "start": 65, "end": 65},
            )
        set_table_geometry(event_table, [8, 84.83, 8, 84.83, 8, 84.84])
        set_table_borders(event_table, color="line", size=5, inside=True)
        for row in event_table.rows:
            prevent_row_split(row, min_height_mm=13.5)

        sequence = doc.add_table(rows=1, cols=12)
        fill_cell(sequence.cell(0, 0), "순서  ·  글자만 써요", size=8.1, bold=True, color="brick", fill="brick_light", align=WD_ALIGN_PARAGRAPH.CENTER)
        widths = [58]
        for col in range(1, 12):
            if col % 2 == 1:
                fill_cell(sequence.cell(0, col), "", fill="white", align=WD_ALIGN_PARAGRAPH.CENTER, margins={"top": 20, "bottom": 20, "start": 20, "end": 20})
                widths.append(18)
            else:
                fill_cell(sequence.cell(0, col), "→", size=8.8, bold=True, color="muted", align=WD_ALIGN_PARAGRAPH.CENTER, margins={"top": 20, "bottom": 20, "start": 15, "end": 15})
                widths.append(10)
        set_table_geometry(sequence, widths, indent_dxa=mm_to_dxa(31))
        set_table_borders(sequence, color="green", size=6, inside=True)
        prevent_row_split(sequence.rows[0], min_height_mm=8.5)

        add_section_heading(
            doc,
            "3  내가 만들 여섯 장면",
            "장면마다 ‘누가·어디·느낌·모습’을 골라 ○표해요. 여러 답이 있어요.",
            compact=True,
        )
        add_scene_planning_grid(doc, story)
        add_chosen_scene_sketch(doc)

    path = OUT / "02_2차시_학생활동지_3종.docx"
    doc.save(path)
    return path


def make_fixed_reading_material() -> Path:
    doc = setup_doc("2차시 고정 이야기 읽기 자료", student=True)
    reading_texts = {
        "해와 달이 된 오누이": [
            "어머니는 아침에 떡을 팔러 장으로 떠났어요.",
            "산길을 가던 어머니는 배고픈 호랑이를 만났어요.",
            "호랑이는 어머니인 척하며 오누이의 집으로 찾아왔어요.",
            "오누이는 호랑이를 피해 밖으로 달아나 큰 나무 위로 올라갔어요.",
            "오누이는 하늘을 향해 자신들을 도와달라고 간절히 빌었어요.",
            "하늘의 도움으로 오누이는 해와 달이 되어 세상을 밝혀 주었어요.",
        ],
        "흥부와 놀부": [
            "흥부와 놀부는 흥부네 집과 놀부네 집에서 서로 다른 마음으로 살았어요.",
            "흥부네 집의 제비집에서 다친 제비를 보고 다리를 정성껏 고쳐 주었어요.",
            "봄이 되자 돌아온 제비는 고마운 마음을 담아 흥부에게 박씨를 물어다 주었어요.",
            "박이 열린 마당에서 흥부 가족이 박을 타자 살림에 큰 도움이 되는 것들이 나왔어요.",
            "놀부네 집에서도 흥부를 따라 하며 박씨를 얻으려고 했어요.",
            "나중에 놀부는 자신의 잘못을 깨닫고, 형제는 서로 도우며 사이좋게 지냈어요.",
        ],
        "의좋은 형제": [
            "형과 동생은 형의 집과 동생의 집 가까운 논에서 각자 농사를 지었어요.",
            "가을이 되자 형제는 거둔 볏단을 똑같이 나누었어요.",
            "형은 동생에게 더 주려고 밤에 볏단을 동생의 집으로 옮겼어요.",
            "동생도 형에게 더 주려고 밤에 볏단을 형의 집으로 옮겼어요.",
            "다음 날에도 볏단 수가 그대로이자 두 사람은 이상하게 생각했어요.",
            "그날 밤, 밤길에서 마주친 형제는 서로의 따뜻한 마음을 알고 꼭 안았어요.",
        ],
    }

    for page_index, story in enumerate(STORIES, start=1):
        if page_index > 1:
            doc.add_page_break()
        add_masthead(
            doc,
            kicker="2차시 · 교사 검토 고정 자료",
            title=story["title"],
            subtitle=f"읽기 자료  ·  {story['theme']}",
            marker=f"읽기 {page_index} / 3",
        )

        add_section_heading(doc, "읽으며 찾아보기", "누가 나오는지, 어디에서 일어나는지, 무슨 일이 생기는지 살펴봐요.")
        focus = doc.add_table(rows=1, cols=3)
        for index, (label, body) in enumerate((
            ("누가", "이야기에 나오는 인물"),
            ("어디서", "일이 일어나는 곳"),
            ("무엇을", "인물이 한 중요한 일"),
        )):
            cell = focus.cell(0, index)
            fill_cell(cell, fill="sage_light", align=WD_ALIGN_PARAGRAPH.CENTER)
            p = cell.paragraphs[0]
            add_text(p, label, size=11.0, bold=True, color="green_dark")
            p2 = style_paragraph(cell.add_paragraph(), before=1, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(p2, body, size=10.2, color="muted")
        set_table_geometry(focus, [62.5, 62.5, 62.5])
        set_table_borders(focus, color="line", size=5, inside=True)
        prevent_row_split(focus.rows[0], min_height_mm=15)

        add_section_heading(doc, "이야기 읽기", "아래 여섯 문장을 천천히 읽어 보세요.")
        passages = doc.add_table(rows=6, cols=2)
        for scene_index, sentence in enumerate(reading_texts[story["title"]], start=1):
            fill_cell(
                passages.cell(scene_index - 1, 0), str(scene_index),
                size=12.0, bold=True, color="white", fill="green",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                margins={"top": 100, "bottom": 100, "start": 60, "end": 60},
            )
            fill_cell(
                passages.cell(scene_index - 1, 1), sentence,
                size=11.0, bold=True, fill="paper" if scene_index % 2 else "white",
                line=1.35,
                margins={"top": 115, "bottom": 115, "start": 160, "end": 160},
            )
            prevent_row_split(passages.rows[scene_index - 1], min_height_mm=22)
        set_table_geometry(passages, [16, 171.5])
        set_table_borders(passages, color="line", size=5, inside=True)

        note = doc.add_table(rows=1, cols=1)
        fill_cell(
            note.cell(0, 0),
            "교사 검토 고정 읽기 자료  ·  생성형 AI 결과물이 아닙니다.",
            size=10.2, bold=True, color="brick", fill="brick_light",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            margins={"top": 70, "bottom": 70, "start": 100, "end": 100},
        )
        set_table_geometry(note, [187.5])
        set_table_borders(note, color="brick", size=5, inside=False)
        prevent_row_split(note.rows[0], min_height_mm=10)

    path = OUT / "05_2차시_고정이야기읽기자료.docx"
    doc.save(path)
    return path


def make_plan() -> Path:
    doc = setup_doc("2차시 교수학습과정안", landscape=True)
    add_masthead(
        doc,
        kicker="AI+교과 내용 융합 교수·학습 과정안",
        title="2차시  우리가 만들 이야기 정하기",
        subtitle="이야기의 뼈대를 이해하고, 학생이 직접 여섯 장면을 기획하는 40분 수업",
        marker="초등 3-4학년 · 40분",
        wide=True,
        compact=True,
    )

    meta = doc.add_table(rows=5, cols=6)
    meta.cell(0, 1).merge(meta.cell(0, 3))
    meta.cell(2, 1).merge(meta.cell(2, 5))
    meta.cell(4, 1).merge(meta.cell(4, 5))
    rows = [
        ["프로그램명", "AI와 함께 만드는 우리 옛이야기 그림책", "", "", "대상", "초등 3-4학년"],
        ["중심 교과", "창체·국어", "차시", "2/10", "수업 시간", "40분"],
        ["성취기준", "[4국05-02] 인물, 사건, 배경에 주목하며 작품을 이해한다.", "", "", "", ""],
        ["AI 영역", "인공지능의 이해", "세부 영역", "인공지능과 사회", "LEAP", "Look"],
        ["AI 내용 요소", "인공지능과의 첫 만남 - 사람과 AI의 인식 방식 차이", "", "", "", ""],
    ]
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = meta.rows[row_index].cells[col_index]
            if not value and cell.text:
                continue
            label = col_index in (0, 2, 4) and bool(value)
            fill_cell(
                cell, value,
                size=7.6 if label else 8.0,
                bold=label,
                color="green_dark" if label else "ink",
                fill="sage" if label else "white",
                align=WD_ALIGN_PARAGRAPH.CENTER if label else WD_ALIGN_PARAGRAPH.LEFT,
                margins={"top": 25, "bottom": 25, "start": 55, "end": 55},
                line=1.05,
            )
        prevent_row_split(meta.rows[row_index], min_height_mm=5.5)
    set_table_geometry(meta, [24, 65, 24, 48, 24, 96])
    set_table_borders(meta, color="line", size=4, inside=True)

    add_section_heading(doc, "학습 목표", "읽기·배열 뒤에 학생 자신의 그림책 장면을 기획합니다.", compact=True)
    goals = doc.add_table(rows=1, cols=3)
    goal_texts = [
        "옛이야기의 인물과 배경을 찾을 수 있다.",
        "중요한 사건 여섯 개를 이야기 순서대로 배열할 수 있다.",
        "각 장면의 인물·장소·느낌·모습을 골라 나만의 여섯 장면을 기획할 수 있다.",
    ]
    for index, goal in enumerate(goal_texts):
        cell = goals.cell(0, index)
        fill_cell(cell, fill="sage_light", margins={"top": 45, "bottom": 45, "start": 70, "end": 70})
        p = cell.paragraphs[0]
        add_text(p, f"목표 {index + 1}  ", size=7.6, bold=True, color="brick")
        add_text(p, goal, size=8.0, bold=True, color="ink")
    set_table_geometry(goals, [93.67, 93.67, 93.66])
    set_table_borders(goals, color="line", size=4, inside=True)
    prevent_row_split(goals.rows[0], min_height_mm=9)

    add_section_heading(doc, "교수·학습 과정", "총 40분", compact=True)
    flow = doc.add_table(rows=7, cols=6)
    headers = ["단계", "시간", "교수 활동", "학생 활동", "자료", "유의점"]
    for index, header in enumerate(headers):
        fill_cell(flow.cell(0, index), header, size=7.9, bold=True, color="green_dark", fill="sage", align=WD_ALIGN_PARAGRAPH.CENTER)
    data = [
        ["도입", "5분", "‘그림책의 한 장면을 그리려면 무엇을 먼저 정해야 할까?’라고 질문한다.", "그림에 필요한 인물·장소·느낌을 자유롭게 말한다.", "PPT 1-3", "학생의 쉬운 낱말을 그대로 받아 준다."],
        ["활동 1", "5분", "교사가 검토한 세 가지 옛이야기의 제목과 짧은 소개를 제시한다.", "모둠에서 만들 이야기 하나를 고르고 까닭을 짧게 말한다.", "PPT 4\n읽기 자료", "한 문장 말하기 또는 손들기 선택으로 참여한다."],
        ["활동 2", "5분", "선택한 이야기를 함께 읽고, 인물과 일이 일어나는 장소를 짚는다.", "이야기에 맞는 인물과 장소에 ○표한다.", "PPT 5-6\n활동지", "이 단계만 이야기 사실을 확인한다. 긴 글쓰기는 하지 않는다."],
        ["활동 3", "10분", "섞인 사건 여섯 개를 읽고 앞뒤 원인을 묻는다.", "사건 카드를 먼저 일어난 일부터 배열하고 순서를 확인한다.", "PPT 7-9\n웹앱", "이야기 순서는 확인하되 점수와 시간 제한은 두지 않는다."],
        ["활동 4", "12분", "장면마다 ‘누가·어디·어떤 느낌·어떤 모습’을 고르는 방법을 한 장면만 시범 보인다.", "여섯 장면의 선택지에 ○표하고 작은 그림으로 장면을 메모한다.", "PPT 10-12\n활동지", "정답을 따라 하게 하지 않는다. 사건을 바꾸지 않는 범위에서 학생 선택을 모두 인정한다."],
        ["정리", "3분", "사건 순서와 장면 기획의 차이를 짚고, 선택 이유를 질문한다.", "가장 마음에 드는 한 장면을 짝에게 보여 주고 선택한 까닭을 말한다.", "PPT 13", "말하기가 어려우면 고른 낱말을 가리켜도 된다."],
    ]
    for row_index, row_data in enumerate(data, start=1):
        for col_index, value in enumerate(row_data):
            fill_cell(
                flow.cell(row_index, col_index), value,
                size=7.8,
                bold=col_index == 0,
                color="green_dark" if col_index == 0 else "ink",
                fill="paper" if row_index % 2 == 0 else "white",
                align=WD_ALIGN_PARAGRAPH.CENTER if col_index in (0, 1, 4) else WD_ALIGN_PARAGRAPH.LEFT,
                margins={"top": 25, "bottom": 25, "start": 50, "end": 50},
                line=1.05,
            )
        prevent_row_split(flow.rows[row_index], min_height_mm=9.5 if row_index in (3, 4, 5) else 8.5)
    set_table_geometry(flow, [17, 14, 74, 74, 34, 68])
    set_table_borders(flow, color="line", size=4, inside=True)

    add_section_heading(doc, "평가와 운영", "교사가 수업 중 관찰합니다.", compact=True)
    notes = doc.add_table(rows=1, cols=3)
    note_data = [
        ("관찰 평가", "이야기 사실과 사건 순서를 확인하는가? 각 장면의 인물·장소·느낌·모습을 스스로 골라 그림책 장면을 기획하는가?"),
        ("준비물", "교사용 PPT, 고정 이야기 읽기 자료, 태블릿 또는 활동지, 필기구"),
        ("수업 원칙", "사실·순서만 확인한다. 장면 기획에는 하나의 정답이 없다. 생성형 AI 없이 교사 검토 고정 자료만 사용한다."),
    ]
    for index, (label, body) in enumerate(note_data):
        cell = notes.cell(0, index)
        fill_cell(cell, fill="brick_light" if index == 2 else "sage_light", margins={"top": 40, "bottom": 40, "start": 70, "end": 70})
        p = cell.paragraphs[0]
        add_text(p, label, size=7.7, bold=True, color="brick" if index == 2 else "green_dark")
        p2 = style_paragraph(cell.add_paragraph(), before=1, line=1.05)
        add_text(p2, body, size=7.5, color="ink")
    set_table_geometry(notes, [93.67, 93.67, 93.66])
    set_table_borders(notes, color="line", size=4, inside=True)
    prevent_row_split(notes.rows[0], min_height_mm=14)

    p = style_paragraph(doc.add_paragraph(), before=1.5, after=0, line=1.0)
    add_text(
        p,
        "※ 원본 교육과정의 2차시 주요 활동을 40분 수업으로 재구성했습니다. 성취기준 번호는 운영 학교 교육과정과 대조 후 확정합니다.",
        size=6.8,
        color="muted",
    )

    path = OUT / "03_2차시_교수학습과정안.docx"
    doc.save(path)
    return path


def make_answer_key() -> Path:
    doc = setup_doc("2차시 교사용 답안", landscape=True)
    add_masthead(
        doc,
        kicker="교사용 자료",
        title="2차시 이야기 확인 및 장면 기획 지도",
        subtitle="사실과 순서만 확인 · 학생의 그림책 기획에는 하나의 정답이 없음",
        marker="TEACHER · LESSON 02",
        wide=True,
    )

    add_section_heading(doc, "확인할 답", "제공한 고정 이야기의 인물·장소·사건 순서만 확인합니다.")
    panels = doc.add_table(rows=1, cols=3)
    for story_index, story in enumerate(STORIES):
        cell = panels.cell(0, story_index)
        fill_cell(cell, fill="paper", margins={"top": 100, "bottom": 100, "start": 120, "end": 120})

        p = style_paragraph(cell.paragraphs[0], after=1, line=1.05, keep_with_next=True)
        add_text(p, story["theme"], size=8.2, bold=True, color="brick")
        p = style_paragraph(cell.add_paragraph(), after=3, line=1.1, keep_with_next=True)
        add_text(p, story["title"], size=13.2, bold=True, color="green_dark")

        correct_characters = " · ".join(label for label, correct in story["characters"] if correct)
        correct_backgrounds = " · ".join(label for label, correct in story["backgrounds"] if correct)
        for label, value in (
            ("인물", correct_characters),
            ("배경", correct_backgrounds),
            ("사건 순서", "B → E → D → A → F → C"),
        ):
            p = style_paragraph(cell.add_paragraph(), after=2, line=1.18, keep_with_next=True)
            add_text(p, f"{label}  ", size=8.7, bold=True, color="green")
            add_text(p, value, size=9.3, bold=label == "사건 순서", color="ink")

        p = style_paragraph(cell.add_paragraph(), before=2, after=1, line=1.05, keep_with_next=True)
        add_text(p, "바른 흐름", size=8.4, bold=True, color="muted")
        for order, event in enumerate(story["events"], start=1):
            p = style_paragraph(cell.add_paragraph(), after=1, line=1.12)
            add_text(p, f"{order}. {event}", size=8.3, color="ink")

        p = style_paragraph(cell.add_paragraph(), before=2, after=0, line=1.1)
        add_text(p, "장면 기획  ", size=8.4, bold=True, color="brick")
        add_text(p, "정답 없음 · 사건과 이어지면 인정", size=8.4, bold=True, color="ink")

    set_table_geometry(panels, [93.67, 93.67, 93.66])
    set_table_borders(panels, color="line", size=5, inside=True)
    prevent_row_split(panels.rows[0], min_height_mm=91)

    add_section_heading(doc, "기획 활동 지도", "쓰기나 그림 실력이 아니라 학생의 선택과 연결을 관찰합니다.")
    guidance = doc.add_table(rows=1, cols=3)
    guidance_data = [
        (
            "정답으로 확인",
            "제공한 글에 나오는 인물과 장소, 사건 여섯 개의 순서만 확인합니다. 사건 순서는 B → E → D → A → F → C입니다.",
        ),
        (
            "정답 없이 인정",
            "장면별 ‘누가·어디·느낌·모습’과 그림 메모는 하나의 답이 없습니다. 중심 사건을 바꾸지 않고 장면과 이어지면 모두 인정합니다.",
        ),
        (
            "초3 구두 발문",
            "“누가 나오게 했니?” “어디에서 일어나니?” “왜 이 느낌을 골랐니?” “그림에서 무엇을 크게 보여 주고 싶니?” 한 문장 또는 가리키기로 답해도 됩니다.",
        ),
    ]
    for index, (label, body) in enumerate(guidance_data):
        cell = guidance.cell(0, index)
        fill_cell(cell, fill="brick_light" if index == 2 else "sage_light", margins={"top": 85, "bottom": 85, "start": 100, "end": 100})
        p = cell.paragraphs[0]
        add_text(p, label, size=9.0, bold=True, color="brick" if index == 2 else "green_dark")
        p2 = style_paragraph(cell.add_paragraph(), before=2, line=1.22)
        add_text(p2, body, size=8.8, color="ink")
    set_table_geometry(guidance, [93.67, 93.67, 93.66])
    set_table_borders(guidance, color="line", size=4, inside=True)
    prevent_row_split(guidance.rows[0], min_height_mm=24)

    p = style_paragraph(doc.add_paragraph(), before=2, after=1, align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0)
    add_text(
        p,
        "운영 원칙: 생성형 AI는 사용하지 않습니다. 학생이 선택지를 고르고 직접 그려 여섯 장면을 기획합니다.",
        size=8.0,
        bold=True,
        color="brick",
    )
    p = style_paragraph(doc.add_paragraph(), after=0, align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.0)
    add_text(p, "판본 유의: 사실과 순서 답은 제공된 고정 자료 기준이며, 다른 판본을 곧바로 오답으로 단정하지 않습니다.", size=8.0, bold=True, color="muted")

    path = OUT / "04_2차시_교사용답안.docx"
    doc.save(path)
    return path


def main():
    outputs = [
        make_worksheet_pack(),
        make_plan(),
        make_answer_key(),
        make_fixed_reading_material(),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
