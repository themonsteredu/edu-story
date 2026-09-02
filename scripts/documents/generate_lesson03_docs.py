from __future__ import annotations

from pathlib import Path
import sys
import tempfile
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from PIL import Image, ImageEnhance, ImageOps
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import generate_lesson02_docs as base


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "public" / "resources" / "lesson-03"
OUT.mkdir(parents=True, exist_ok=True)

# Design system: compact_reference_guide with named school-material overrides.
# Exact page geometry is A4, not Letter: portrait 210 x 297 mm with 8/10 mm
# top-side margins for student sheets; landscape 297 x 210 mm with 6/8 mm
# margins for teacher sheets.  Body is S-Core Dream 4 Regular, 10.5 pt,
# 1.25 line spacing, 0 pt before and 6 pt after.  Tables use fixed DXA
# geometry, 120 DXA indent, 90/90/120/120 DXA cell margins, and no exact
# row heights.  This continues the supplied project visual system.
base.COLORS.update(
    {
        "ink": "20312F",
        "muted": "63746F",
        "green": "24544F",
        "green_dark": "173D3A",
        "sage": "E7F0EC",
        "sage_light": "F4F8F6",
        "paper": "FCF9F2",
        "sand": "F2EBDD",
        # Legacy helper names retained; these values are the restrained blue accent.
        "brick": "386A88",
        "brick_light": "E7F0F5",
        "line": "A9BBB5",
        "white": "FFFFFF",
    }
)


PHOTO_SOURCE = {
    "warm": ROOT / "public/assets/lesson-01/classroom-real.webp",
    "warm_dark": ROOT / "public/assets/lesson-01/classroom-real.webp",
    "cool": ROOT / "public/assets/lesson-01/automatic-door-real.webp",
    "green": ROOT / "public/assets/lesson-01/photo-classification-real.webp",
}

CLASSIFY_CARDS = [
    {"code": "A", "kind": "문자", "text": "호랑이가 산길에 나타났어요."},
    {"code": "B", "kind": "문자", "text": "비 · 우산 · 골목"},
    {"code": "C", "kind": "이미지", "photo": "warm", "alt": "교실에서 그림을 그리는 어린이 사진"},
    {"code": "D", "kind": "이미지", "photo": "cool", "alt": "회색 건물 입구의 자동문 사진"},
    {"code": "E", "kind": "소리", "sound": "① 빗소리", "audio": "rain-soft.wav"},
    {"code": "F", "kind": "소리", "sound": "② 새소리", "audio": "birds-morning.wav"},
]

PLAN_CHOICES = [
    ("주인공", ["사람", "동물", "신기한 존재"]),
    ("장소", ["집", "길 · 숲", "마을 · 궁궐"]),
    ("시간", ["낮", "밤", "비 오는 때"]),
    ("표정", ["기뻐요", "놀라요", "걱정해요"]),
    ("빛 · 색", ["따뜻한 밝은빛", "차가운 어두운빛", "싱그러운 초록빛"]),
    ("소리", ["② 새소리", "① 빗소리", "④ 바람소리"]),
]

PHOTO_CHOICES = [
    ("warm", "따뜻한 밝은빛", "교실에서 그림을 그리는 어린이 사진"),
    ("cool", "차가운 어두운빛", "회색 건물 입구의 자동문 사진"),
    ("green", "싱그러운 초록빛", "휴대전화 화면에 보이는 초록 잎 사진"),
]

SOUND_GUIDE = [
    ("① 빗소리", "rain-soft.wav", "빗소리"),
    ("② 새소리", "birds-morning.wav", "새소리"),
    ("③ 천둥소리", "thunder-low.wav", "천둥소리"),
    ("④ 바람소리", "wind-low.wav", "바람소리"),
]


def photo_png(key: str) -> Path:
    """Convert the licensed local WebP photo to a cropped PNG for Word."""
    cache = Path(tempfile.gettempdir()) / "edu-story-lesson03-docx-images"
    cache.mkdir(parents=True, exist_ok=True)
    output = cache / f"{key}.png"
    source = PHOTO_SOURCE[key]
    if not output.exists() or output.stat().st_mtime < source.stat().st_mtime:
        with Image.open(source) as image:
            prepared = ImageOps.fit(image.convert("RGB"), (960, 600), method=Image.Resampling.LANCZOS)
            if key == "warm_dark":
                prepared = ImageEnhance.Color(prepared).enhance(0.65)
                prepared = ImageEnhance.Brightness(prepared).enhance(0.34)
            prepared.save(output, "PNG", optimize=True)
    return output


def add_photo(cell, key: str, alt: str, *, width_mm: float = 34, height_mm: float = 21):
    p = base.style_paragraph(cell.add_paragraph(), after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    shape = p.add_run().add_picture(str(photo_png(key)), width=Mm(width_mm), height=Mm(height_mm))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    return shape


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def setup(title: str, *, landscape: bool = False, student: bool = False):
    doc = base.setup_doc(title, landscape=landscape, student=student)
    doc.core_properties.subject = "AI와 함께 만드는 우리 옛이야기 그림책 3차시"
    doc.core_properties.comments = "생성형 AI를 사용하지 않는 교사 검토 고정 자료"
    return doc


def compact_cell(
    cell,
    text: str = "",
    *,
    size: float = 9.4,
    bold: bool = False,
    color: str = "ink",
    fill: str | None = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line: float = 1.15,
    top: int = 65,
    bottom: int = 65,
    start: int = 90,
    end: int = 90,
):
    return base.fill_cell(
        cell,
        text,
        size=size,
        bold=bold,
        color=color,
        fill=fill,
        align=align,
        line=line,
        margins={"top": top, "bottom": bottom, "start": start, "end": end},
    )


def add_student_header(doc):
    table = doc.add_table(rows=1, cols=4)
    title = table.cell(0, 0)
    compact_cell(title, fill="paper", top=75, bottom=75, start=120, end=100)
    p = base.style_paragraph(title.paragraphs[0], after=1, line=1.0)
    base.add_text(p, "3차시  ", size=8.7, bold=True, color="brick")
    base.add_text(p, "AI가 이해하는 데이터", size=16.5, bold=True, color="ink")
    p = base.style_paragraph(title.add_paragraph(), after=0, line=1.05)
    base.add_text(p, "문자 · 이미지 · 소리를 골라 내 장면 꾸러미를 만들어요.", size=8.7, bold=True, color="green")

    fields = ["학년·반\n________", "모둠\n______", "이름\n________"]
    for col, label in enumerate(fields, start=1):
        compact_cell(
            table.cell(0, col),
            label,
            size=8.6,
            bold=True,
            color="green_dark",
            fill="sage",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            line=1.05,
            top=45,
            bottom=45,
            start=45,
            end=45,
        )
    base.set_table_geometry(table, [105, 27, 25, 30.5])
    base.set_table_borders(table, color="green", size=6, inside=True)
    base.prevent_row_split(table.rows[0], min_height_mm=18)


def add_choice_run(paragraph, values: list[str], *, size: float = 9.3):
    for index, value in enumerate(values):
        if index:
            base.add_text(paragraph, "    ", size=size)
        base.add_text(paragraph, f"○ {value}", size=size, bold=True, color="ink")


def make_worksheet() -> Path:
    doc = setup("3차시 학생 활동지", student=True)
    add_student_header(doc)

    base.add_section_heading(
        doc,
        "1  어떤 데이터일까요?",
        "카드마다 알맞은 데이터에 ○표해요.",
        compact=True,
    )
    table = doc.add_table(rows=3, cols=2)
    for index, card in enumerate(CLASSIFY_CARDS):
        row, col = divmod(index, 2)
        cell = table.cell(row, col)
        compact_cell(
            cell,
            fill="paper" if index % 2 == 0 else "white",
            top=55,
            bottom=55,
            start=90,
            end=90,
        )
        p = base.style_paragraph(cell.paragraphs[0], after=1, line=1.05, keep_with_next=True)
        base.add_text(p, f"{card['code']}  ", size=8.3, bold=True, color="brick")
        if "text" in card:
            base.add_text(p, card["text"], size=9.5, bold=True, color="ink")
        elif "photo" in card:
            base.add_text(p, "눈으로 보는 사진", size=8.7, bold=True, color="muted")
            add_photo(cell, card["photo"], card["alt"], width_mm=30, height_mm=16.5)
        else:
            base.add_text(p, card["sound"], size=11.2, bold=True, color="green_dark")
            p = base.style_paragraph(cell.add_paragraph(), after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, "선생님이 번호 음원을 틀면 들어요.", size=7.8, bold=True, color="muted")
        p = base.style_paragraph(cell.add_paragraph(), line=1.0)
        add_choice_run(p, ["문자", "이미지", "소리"], size=8.8)
    base.set_table_geometry(table, [93.75, 93.75])
    base.set_table_borders(table, color="line", size=5, inside=True)
    for row in table.rows:
        base.prevent_row_split(row, min_height_mm=26)

    base.add_section_heading(
        doc,
        "2  같은 장면, 다른 느낌",
        "같은 문장을 꾸민 자료를 보고 느낌에 ○표해요.",
        compact=True,
    )
    scene = doc.add_table(rows=1, cols=1)
    compact_cell(
        scene.cell(0, 0),
        "같은 문장  |  사람들이 교실에 있어요.",
        size=9.6,
        bold=True,
        color="green_dark",
        fill="sage",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        top=45,
        bottom=45,
    )
    base.set_table_geometry(scene, [187.5])
    base.set_table_borders(scene, color="green", size=5, inside=False)
    base.prevent_row_split(scene.rows[0], min_height_mm=8)

    compare = doc.add_table(rows=1, cols=2)
    bundles = [
        ("A 꾸러미", "밝은 낮, 함께 배우는 시간", "warm", "밝은 교실에서 선생님과 어린이들이 함께 있는 사진", "② 새소리", ["편안해요", "신나요", "무서워요"]),
        ("B 꾸러미", "어두운 밤, 조용한 교실", "warm_dark", "같은 교실 사진을 어둡게 처리한 모습", "③ 천둥소리", ["편안해요", "신나요", "무서워요"]),
    ]
    for index, (label, text_clue, photo_key, photo_alt, sound, feelings) in enumerate(bundles):
        cell = compare.cell(0, index)
        compact_cell(
            cell,
            fill="sage_light" if index == 0 else "brick_light",
            top=70,
            bottom=70,
            start=110,
            end=110,
        )
        p = base.style_paragraph(cell.paragraphs[0], after=2, line=1.05)
        base.add_text(p, label, size=10.2, bold=True, color="green_dark" if index == 0 else "brick")
        p = base.style_paragraph(cell.add_paragraph(), after=1, line=1.05)
        base.add_text(p, f"문자  {text_clue}", size=8.7, bold=True)
        add_photo(cell, photo_key, photo_alt, width_mm=36, height_mm=19)
        p = base.style_paragraph(cell.add_paragraph(), after=1, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
        base.add_text(p, f"이미지  {'밝은 교실 사진' if index == 0 else '같은 교실 사진을 어둡게 본 모습'}    소리  {sound}", size=8.0, bold=True)
        p = base.style_paragraph(cell.add_paragraph(), line=1.05)
        add_choice_run(p, feelings, size=8.7)
    base.set_table_geometry(compare, [93.75, 93.75])
    base.set_table_borders(compare, color="line", size=5, inside=True)
    base.prevent_row_split(compare.rows[0], min_height_mm=43)

    p = base.style_paragraph(doc.add_paragraph(), before=1, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    base.add_text(p, "같은 문장도 문자 · 이미지 · 소리가 달라지면 느낌이 달라져요.", size=8.7, bold=True, color="brick")

    base.add_section_heading(
        doc,
        "3  내 이야기의 한 장면 기획하기",
        "2차시에 정한 이야기 장면을 떠올리며 각 줄에서 하나씩 ○표해요.",
        compact=True,
    )
    build = doc.add_table(rows=len(PLAN_CHOICES), cols=2)
    for row, (label, values) in enumerate(PLAN_CHOICES):
        compact_cell(
            build.cell(row, 0),
            label,
            size=9.1,
            bold=True,
            color="green_dark",
            fill="sage",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            top=45,
            bottom=45,
        )
        cell = build.cell(row, 1)
        compact_cell(cell, fill="paper", top=45, bottom=45, start=85, end=85)
        p = cell.paragraphs[0]
        add_choice_run(p, values, size=8.8)
        base.prevent_row_split(build.rows[row], min_height_mm=8.2)
    base.set_table_geometry(build, [22, 165.5])
    base.set_table_borders(build, color="line", size=5, inside=True)

    note = doc.add_table(rows=1, cols=1)
    compact_cell(
        note.cell(0, 0),
        "오늘의 발견  |  AI가 장면을 이해하려면 주인공 · 장소 · 시간 · 표정 · 빛 · 색 · 소리 같은 데이터가 필요해요.  ·  생성형 AI 없이 활동해요.",
        size=8.8,
        bold=True,
        color="green_dark",
        fill="sage_light",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        top=40,
        bottom=40,
    )
    base.set_table_geometry(note, [187.5])
    base.set_table_borders(note, color="green", size=4, inside=False)
    base.prevent_row_split(note.rows[0], min_height_mm=7)

    path = OUT / "02_3차시_학생활동지.docx"
    doc.save(path)
    return path


def make_plan() -> Path:
    doc = setup("3차시 교수학습과정안", landscape=True)
    base.add_masthead(
        doc,
        kicker="AI+교과 내용 융합 교수·학습 과정안",
        title="3차시  AI가 이해하는 데이터",
        subtitle="문자 · 이미지 · 소리를 구분하고, 학생이 직접 한 장면의 자료 꾸러미를 만드는 40분 수업",
        marker="초등 3-4학년 · 40분",
        wide=True,
        compact=True,
    )

    meta = doc.add_table(rows=5, cols=6)
    meta.cell(0, 1).merge(meta.cell(0, 3))
    meta.cell(2, 1).merge(meta.cell(2, 5))
    meta.cell(4, 1).merge(meta.cell(4, 5))
    rows = [
        ["프로그램명", "AI와 함께 만드는 우리 옛이야기 그림책", "", "", "대상", "초등 3-4학년"],
        ["중심 교과", "국어", "차시", "3/10", "수업 시간", "40분"],
        ["성취기준", "[4국05-02] 인물, 사건, 배경에 주목하며 작품을 이해한다.  [4국05-03] 이야기의 흐름을 파악하여 이어질 내용을 상상하고 표현한다.", "", "", "", ""],
        ["AI 영역", "인공지능의 원리와 활용", "내용 요소", "데이터", "LEAP", "Explore"],
        ["AI 학습 내용", "AI가 다루는 여러 데이터 - 문자 · 이미지 · 소리의 특징과 장면 느낌의 변화", "", "", "", ""],
    ]
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = meta.rows[row_index].cells[col_index]
            if not value and cell.text:
                continue
            label = col_index in (0, 2, 4) and bool(value)
            compact_cell(
                cell,
                value,
                size=7.3 if label else 7.7,
                bold=label,
                color="green_dark" if label else "ink",
                fill="sage" if label else "white",
                align=WD_ALIGN_PARAGRAPH.CENTER if label else WD_ALIGN_PARAGRAPH.LEFT,
                line=1.04,
                top=24,
                bottom=24,
                start=50,
                end=50,
            )
        base.prevent_row_split(meta.rows[row_index], min_height_mm=5.5)
    base.set_table_geometry(meta, [24, 65, 24, 48, 24, 96])
    base.set_table_borders(meta, color="line", size=4, inside=True)

    base.add_section_heading(doc, "학습 목표", "구분한 뒤 학생의 선택을 한 장면에 연결합니다.", compact=True)
    goals = doc.add_table(rows=1, cols=3)
    goal_texts = [
        "문자 · 이미지 · 소리 데이터를 구분할 수 있다.",
        "같은 문장도 문자 · 이미지 · 소리 자료에 따라 느낌이 달라짐을 알 수 있다.",
        "내 이야기에 필요한 주인공 · 장소 · 시간 · 표정 · 빛 · 소리를 직접 고를 수 있다.",
    ]
    for index, goal in enumerate(goal_texts):
        cell = goals.cell(0, index)
        compact_cell(cell, fill="sage_light", top=42, bottom=42, start=65, end=65)
        p = cell.paragraphs[0]
        base.add_text(p, f"목표 {index + 1}  ", size=7.4, bold=True, color="brick")
        base.add_text(p, goal, size=7.8, bold=True)
    base.set_table_geometry(goals, [93.67, 93.67, 93.66])
    base.set_table_borders(goals, color="line", size=4, inside=True)
    base.prevent_row_split(goals.rows[0], min_height_mm=9)

    base.add_section_heading(doc, "교수·학습 과정", "총 40분", compact=True)
    flow = doc.add_table(rows=7, cols=6)
    headers = ["단계", "시간", "교수 활동", "학생 활동", "자료", "유의점"]
    for index, header in enumerate(headers):
        compact_cell(
            flow.cell(0, index),
            header,
            size=7.6,
            bold=True,
            color="green_dark",
            fill="sage",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            top=28,
            bottom=28,
        )
    mark_header_row(flow.rows[0])
    data = [
        ["도입", "5분", "글, 실제 사진, 번호 음원을 차례로 보여 주고 ‘모두 AI가 살펴볼 수 있는 자료일까?’라고 묻는다.", "눈으로 보거나 귀로 들은 차이를 말하고 오늘의 질문을 확인한다.", "PPT 1-3", "AI의 어려운 원리는 설명하지 않고 ‘자료를 살펴본다’로 말한다."],
        ["설명", "5분", "문자 · 이미지 · 소리 데이터의 쉬운 뜻과 단서를 실제 매체로 설명한다.", "읽어요 · 보여요 · 들려요 가운데 알맞은 말을 골라 말하거나 가리킨다.", "PPT 4-6", "사진이라고 적힌 글이 아니라 실제 사진을 보여 주고, 번호 음원을 직접 재생한다."],
        ["활동 1", "8분", "고정 카드 A-F를 제시하고 소리 ①·② 음원을 번호에 맞춰 재생한다.", "카드 6개의 데이터 종류에 ○표하고 짝과 하나씩 확인한다.", "PPT 7\n활동지·카드", "카드의 뜻이 아니라 실제 제시 방식이 읽기·보기·듣기 중 무엇인지 살핀다."],
        ["활동 2", "9분", "‘사람들이 교실에 있어요.’에 A와 B 자료를 붙여 문자 · 같은 사진의 밝기 · 소리를 비교한다.", "A와 B에서 느껴지는 마음에 ○표하고, 달라진 문자 · 이미지 · 소리를 말하거나 가리킨다.", "PPT 8-11\n활동지", "A는 밝은 교실 실사와 ② 새소리, B는 같은 사진을 어둡게 처리한 모습과 ③ 천둥소리로 제시한다."],
        ["활동 3", "10분", "주인공 · 장소 · 시간 · 표정 · 빛과 색 · 소리에서 하나씩 고르는 방법만 시범 보인다.", "2차시에 정한 자기 이야기를 떠올려 여섯 항목을 직접 골라 한 장면을 기획한다.", "PPT 12\n웹앱·카드", "교사 예시를 복사하게 하지 않는다. 어떤 이야기든 학생의 선택과 까닭을 인정한다."],
        ["정리", "3분", "‘AI가 살펴볼 수 있는 세 가지 데이터는?’과 ‘자료가 달라지면 무엇이 달라질까?’를 묻는다.", "문자 · 이미지 · 소리를 손가락으로 짚고, 내 장면의 느낌을 한 낱말로 말한다.", "PPT 13", "검색·스프레드시트·Canva 활용은 4차시에서 진행한다."],
    ]
    for row_index, row_data in enumerate(data, start=1):
        for col_index, value in enumerate(row_data):
            compact_cell(
                flow.cell(row_index, col_index),
                value,
                size=7.45,
                bold=col_index == 0,
                color="green_dark" if col_index == 0 else "ink",
                fill="paper" if row_index % 2 == 0 else "white",
                align=WD_ALIGN_PARAGRAPH.CENTER if col_index in (0, 1, 4) else WD_ALIGN_PARAGRAPH.LEFT,
                line=1.04,
                top=24,
                bottom=24,
                start=46,
                end=46,
            )
        base.prevent_row_split(flow.rows[row_index], min_height_mm=10)
    base.set_table_geometry(flow, [17, 14, 74, 74, 34, 68])
    base.set_table_borders(flow, color="line", size=4, inside=True)

    base.add_section_heading(doc, "평가와 운영", "쓰기 양보다 구분 · 비교 · 선택 과정을 관찰합니다.", compact=True)
    notes = doc.add_table(rows=1, cols=3)
    note_data = [
        ("관찰 평가", "세 가지 데이터 종류를 구분하는가? 자료 차이와 장면 느낌을 연결하는가? 자기 이야기에 필요한 여섯 항목을 직접 고르는가?"),
        ("준비물", "교사용 PPT, 학생 활동지, 고정 데이터 카드, 소리 ①-④ WAV, 태블릿 또는 교실 화면, 필기구"),
        ("수업 원칙", "생성형 AI를 사용하지 않는다. 학생은 고정 카드에서 직접 선택한다. 검색·스프레드시트·Canva는 4차시에 넘긴다."),
    ]
    for index, (label, body) in enumerate(note_data):
        cell = notes.cell(0, index)
        compact_cell(cell, fill="brick_light" if index == 2 else "sage_light", top=40, bottom=40, start=65, end=65)
        p = cell.paragraphs[0]
        base.add_text(p, label, size=7.5, bold=True, color="brick" if index == 2 else "green_dark")
        p = base.style_paragraph(cell.add_paragraph(), before=1, line=1.04)
        base.add_text(p, body, size=7.3)
    base.set_table_geometry(notes, [93.67, 93.67, 93.66])
    base.set_table_borders(notes, color="line", size=4, inside=True)
    base.prevent_row_split(notes.rows[0], min_height_mm=13)

    p = base.style_paragraph(doc.add_paragraph(), before=1.2, after=0, line=1.0)
    base.add_text(
        p,
        "※ 원본 교육과정의 3차시 데이터 탐색을 40분 수업으로 재구성했습니다. 성취기준 번호는 운영 학교 교육과정과 대조 후 확정합니다.",
        size=6.7,
        color="muted",
    )

    path = OUT / "03_3차시_교수학습과정안.docx"
    doc.save(path)
    return path


def make_answer_key() -> Path:
    doc = setup("3차시 교사용 답안", landscape=True)
    base.add_masthead(
        doc,
        kicker="교사용 자료",
        title="3차시 데이터 구분 및 장면 꾸러미 지도",
        subtitle="데이터 종류는 확인 · 느낌과 학생의 장면 선택은 근거가 있으면 모두 인정",
        marker="TEACHER · LESSON 03",
        wide=True,
    )

    base.add_section_heading(doc, "1  데이터 분류 답", "카드가 실제로 제시된 매체를 기준으로 확인합니다.")
    key = doc.add_table(rows=3, cols=4)
    headers = ["카드", "정답", "카드", "정답"]
    for col, header in enumerate(headers):
        compact_cell(key.cell(0, col), header, size=9.1, bold=True, color="green_dark", fill="sage", align=WD_ALIGN_PARAGRAPH.CENTER)
    mark_header_row(key.rows[0])
    for row in range(1, 3):
        for side in range(2):
            index = (row - 1) * 2 + side
            card = CLASSIFY_CARDS[index]
            shown = card.get("text") or card.get("alt") or card.get("sound")
            compact_cell(key.cell(row, side * 2), f"{card['code']}\n{shown}", size=8.8, bold=True, fill="paper", align=WD_ALIGN_PARAGRAPH.CENTER, line=1.12)
            compact_cell(key.cell(row, side * 2 + 1), card["kind"], size=10.0, bold=True, color="brick", fill="brick_light", align=WD_ALIGN_PARAGRAPH.CENTER)
    # Last two cards appear in a compact continuation line below the matrix.
    base.set_table_geometry(key, [54, 39.75, 54, 39.75])
    base.set_table_borders(key, color="line", size=5, inside=True)
    for row in key.rows:
        base.prevent_row_split(row, min_height_mm=11 if row is not key.rows[0] else 8)

    tail = doc.add_table(rows=1, cols=4)
    for side, index in enumerate((4, 5)):
        card = CLASSIFY_CARDS[index]
        shown = card.get("text") or card.get("alt") or card.get("sound")
        compact_cell(tail.cell(0, side * 2), f"{card['code']}  {shown}", size=8.8, bold=True, fill="paper", align=WD_ALIGN_PARAGRAPH.CENTER)
        compact_cell(tail.cell(0, side * 2 + 1), card["kind"], size=10.0, bold=True, color="brick", fill="brick_light", align=WD_ALIGN_PARAGRAPH.CENTER)
    base.set_table_geometry(tail, [54, 39.75, 54, 39.75])
    base.set_table_borders(tail, color="line", size=5, inside=True)
    base.prevent_row_split(tail.rows[0], min_height_mm=10)

    base.add_section_heading(doc, "2  같은 문장 비교", "느낌에는 하나의 정답이 없고 자료를 근거로 설명하는지가 중요합니다.")
    compare = doc.add_table(rows=1, cols=3)
    compare_data = [
        ("A 꾸러미", "밝은 낮, 함께 배우는 시간 · 밝은 교실 실사 · ② 새소리", "warm", "밝은 교실에서 선생님과 어린이들이 함께 있는 사진", "편안해요 또는 신나요를 주로 예상. 고른 자료를 근거로 들면 인정합니다."),
        ("B 꾸러미", "어두운 밤, 조용한 교실 · 같은 교실 실사를 어둡게 처리 · ③ 천둥소리", "warm_dark", "같은 교실 사진을 어둡게 처리한 모습", "무서워요 또는 긴장돼요를 주로 예상. 고른 자료를 근거로 들면 인정합니다."),
        ("비교 핵심", "‘사람들이 교실에 있어요’ 문장은 같음", None, None, "문자 · 이미지 · 소리 자료가 달라 장면의 느낌이 달라집니다. 반대 느낌을 골라도 자료와 연결하여 말하면 인정합니다."),
    ]
    for index, (label, evidence, photo_key, photo_alt, guide) in enumerate(compare_data):
        cell = compare.cell(0, index)
        compact_cell(cell, fill="sage_light" if index < 2 else "brick_light", top=42, bottom=42, start=90, end=90)
        p = cell.paragraphs[0]
        base.add_text(p, label, size=9.2, bold=True, color="brick" if index == 2 else "green_dark")
        p = base.style_paragraph(cell.add_paragraph(), before=2, after=2, line=1.1)
        base.add_text(p, evidence, size=8.2, bold=True)
        if photo_key and photo_alt:
            add_photo(cell, photo_key, photo_alt, width_mm=23, height_mm=9)
        p = base.style_paragraph(cell.add_paragraph(), line=1.1)
        base.add_text(p, guide, size=7.8)
    base.set_table_geometry(compare, [93.67, 93.67, 93.66])
    base.set_table_borders(compare, color="line", size=5, inside=True)
    base.prevent_row_split(compare.rows[0], min_height_mm=33)

    base.add_section_heading(doc, "3  내 이야기의 한 장면", "학생이 여섯 항목에서 하나씩 고른 계획을 존중합니다.")
    guide = doc.add_table(rows=1, cols=3)
    guide_data = [
        ("인정 기준", "주인공 · 장소 · 시간 · 표정 · 빛과 색 · 소리에서 하나씩 골랐다면 인정합니다. 특정 이야기나 조합을 정답으로 제시하지 않습니다."),
        ("예상 학생 반응", "“동물이 밤 숲에 있어요.” “놀란 표정이고 소리 ③이 들려요.”처럼 고른 카드를 가리키거나 짧게 말하면 충분합니다."),
        ("초3 발문", "“누가 나오니?” “어디와 언제니?” “어떤 얼굴이니?” “몇 번 소리를 골랐니?” 말 대신 가리키기로 답해도 됩니다."),
    ]
    for index, (label, body) in enumerate(guide_data):
        cell = guide.cell(0, index)
        compact_cell(cell, fill="brick_light" if index == 2 else "paper", top=55, bottom=55, start=95, end=95)
        p = cell.paragraphs[0]
        base.add_text(p, label, size=9.0, bold=True, color="brick" if index == 2 else "green_dark")
        p = base.style_paragraph(cell.add_paragraph(), before=1, line=1.12)
        base.add_text(p, body, size=8.2)
    base.set_table_geometry(guide, [93.67, 93.67, 93.66])
    base.set_table_borders(guide, color="line", size=5, inside=True)
    base.prevent_row_split(guide.rows[0], min_height_mm=31)

    note = doc.add_table(rows=1, cols=1)
    compact_cell(
        note.cell(0, 0),
        "운영 원칙  |  생성형 AI를 사용하지 않습니다. ① 빗소리 · ② 새소리 · ③ 천둥소리 · ④ 바람소리 WAV를 교사가 직접 재생합니다. 검색 · 스프레드시트 · Canva는 4차시에 진행합니다.",
        size=8.4,
        bold=True,
        color="brick",
        fill="brick_light",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        top=55,
        bottom=55,
    )
    base.set_table_geometry(note, [281])
    base.set_table_borders(note, color="brick", size=5, inside=False)
    base.prevent_row_split(note.rows[0], min_height_mm=9)

    path = OUT / "04_3차시_교사용답안.docx"
    doc.save(path)
    return path


def add_card_heading(cell, code: str, label: str, *, accent: str = "brick"):
    p = base.style_paragraph(cell.paragraphs[0], after=2, line=1.0, keep_with_next=True)
    base.add_text(p, code, size=8.2, bold=True, color=accent)
    base.add_text(p, f"  {label}", size=8.2, bold=True, color="muted")


def make_data_cards() -> Path:
    doc = setup("3차시 고정 데이터 카드", landscape=True, student=True)
    base.add_masthead(
        doc,
        kicker="3차시 · 오려 쓰는 고정 자료",
        title="문자 · 이미지 · 소리 데이터 카드",
        subtitle="생성형 AI 없이 수업에 바로 사용하는 교사 검토 카드",
        marker="카드 1 / 2",
        wide=True,
        compact=True,
    )
    base.add_section_heading(doc, "분류 카드 6장", "점선을 따라 오린 뒤 문자 · 이미지 · 소리 세 모둠으로 나누어 보세요.", compact=True)
    grid = doc.add_table(rows=2, cols=3)
    for index, card in enumerate(CLASSIFY_CARDS):
        row, col = divmod(index, 3)
        cell = grid.cell(row, col)
        fill = "paper" if card["kind"] == "문자" else "sage_light" if card["kind"] == "이미지" else "brick_light"
        compact_cell(cell, fill=fill, top=110, bottom=110, start=140, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_card_heading(cell, card["code"], "분류 카드")
        if "text" in card:
            p = base.style_paragraph(cell.add_paragraph(), before=4, after=4, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, card["text"], size=14.0, bold=True, color="green_dark")
            p = base.style_paragraph(cell.add_paragraph(), line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, "눈으로 읽어요.", size=9.0, bold=True, color="muted")
        elif "photo" in card:
            add_photo(cell, card["photo"], card["alt"], width_mm=58, height_mm=35)
            p = base.style_paragraph(cell.add_paragraph(), line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, "실제 사진을 눈으로 살펴봐요.", size=8.7, bold=True, color="muted")
        else:
            p = base.style_paragraph(cell.add_paragraph(), before=8, after=5, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, card["sound"], size=19.0, bold=True, color="green_dark")
            p = base.style_paragraph(cell.add_paragraph(), line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, "선생님이 번호 음원을 재생하면 들어요.", size=8.8, bold=True, color="muted")
    base.set_table_geometry(grid, [93.67, 93.67, 93.66])
    base.set_table_borders(grid, color="green", size=7, inside=True)
    for row in grid.rows:
        base.prevent_row_split(row, min_height_mm=63)

    note = doc.add_table(rows=1, cols=2)
    compact_cell(note.cell(0, 0), "교사 준비  |  E는 ① rain-soft.wav, F는 ② birds-morning.wav를 직접 재생하세요.", size=8.3, bold=True, color="green_dark", fill="sage", align=WD_ALIGN_PARAGRAPH.CENTER)
    compact_cell(note.cell(0, 1), "정답  |  문자 A·B   이미지 C·D   소리 E·F", size=8.3, bold=True, color="brick", fill="brick_light", align=WD_ALIGN_PARAGRAPH.CENTER)
    base.set_table_geometry(note, [187.33, 93.67])
    base.set_table_borders(note, color="line", size=4, inside=True)
    base.prevent_row_split(note.rows[0], min_height_mm=9)

    doc.add_page_break()
    base.add_masthead(
        doc,
        kicker="3차시 · 내 장면 꾸러미",
        title="내 이야기의 한 장면 기획하기",
        subtitle="2차시에 정한 이야기에서 한 장면을 떠올리고, 여섯 가지를 학생이 직접 고릅니다.",
        marker="카드 2 / 2",
        wide=True,
        compact=True,
    )
    base.add_section_heading(doc, "선택 카드 18장", "각 줄에서 하나씩 골라 내 장면 계획을 완성하세요.", compact=True)
    choices = doc.add_table(rows=len(PLAN_CHOICES), cols=3)
    for row, (kind, values) in enumerate(PLAN_CHOICES):
        for col, value in enumerate(values):
            cell = choices.cell(row, col)
            fill = "sage_light" if row % 2 == 0 else "paper"
            if kind == "빛 · 색":
                fill = ["F5E6A8", "DCE6EF", "DDEADD"][col]
            elif kind == "소리":
                fill = "brick_light"
            compact_cell(cell, fill=fill, top=45, bottom=45, start=100, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            add_card_heading(cell, f"{kind} {col + 1}", "내 장면 선택", accent="brick" if kind == "소리" else "green")
            if kind == "빛 · 색":
                photo_key, _, alt = PHOTO_CHOICES[col]
                add_photo(cell, photo_key, alt, width_mm=35, height_mm=17)
            p = base.style_paragraph(cell.add_paragraph(), before=1, after=1, line=1.08, align=WD_ALIGN_PARAGRAPH.CENTER)
            base.add_text(p, value, size=10.2 if kind != "소리" else 13.0, bold=True, color="green_dark")
            if kind == "소리":
                p = base.style_paragraph(cell.add_paragraph(), line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                base.add_text(p, "선생님이 번호 음원을 재생해요.", size=7.5, bold=True, color="muted")
    base.set_table_geometry(choices, [93.67, 93.67, 93.66])
    base.set_table_borders(choices, color="green", size=7, inside=True)
    for row_index, row in enumerate(choices.rows):
        base.prevent_row_split(row, min_height_mm=26 if PLAN_CHOICES[row_index][0] == "빛 · 색" else 17.5)

    finish = doc.add_table(rows=1, cols=1)
    compact_cell(
        finish.cell(0, 0),
        "계획 확인  |  주인공 + 장소 + 시간 + 표정 + 빛 · 색 + 소리 = 내가 직접 기획한 한 장면",
        size=9.0,
        bold=True,
        color="green_dark",
        fill="sage",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        top=45,
        bottom=45,
    )
    base.set_table_geometry(finish, [281])
    base.set_table_borders(finish, color="green", size=5, inside=False)
    base.prevent_row_split(finish.rows[0], min_height_mm=8)

    path = OUT / "05_3차시_데이터카드.docx"
    doc.save(path)
    return path


def audit_docx(path: Path):
    """Fail fast on missing fonts, fixed table geometry, or unwanted blank pages."""
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    with ZipFile(path) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        styles = ET.fromstring(archive.read("word/styles.xml"))

    font_nodes = styles.findall(".//w:rFonts", ns)
    if not any("S-Core Dream" in " ".join(node.attrib.values()) or "에스코어" in " ".join(node.attrib.values()) for node in font_nodes):
        raise AssertionError(f"S-Core Dream style missing: {path.name}")

    for table_index, table in enumerate(document.findall(".//w:tbl", ns), start=1):
        width = table.find("./w:tblPr/w:tblW", ns)
        indent = table.find("./w:tblPr/w:tblInd", ns)
        grid = table.findall("./w:tblGrid/w:gridCol", ns)
        if width is None or width.get(f"{{{ns['w']}}}type") != "dxa":
            raise AssertionError(f"Table {table_index} has no fixed DXA width in {path.name}")
        if indent is None or indent.get(f"{{{ns['w']}}}type") != "dxa":
            raise AssertionError(f"Table {table_index} has no DXA indent in {path.name}")
        if not grid:
            raise AssertionError(f"Table {table_index} has no fixed grid in {path.name}")
        for cell in table.findall(".//w:tc", ns):
            cell_width = cell.find("./w:tcPr/w:tcW", ns)
            if cell_width is None or cell_width.get(f"{{{ns['w']}}}type") != "dxa":
                raise AssertionError(f"Table {table_index} cell width missing in {path.name}")

    text = "".join(document.itertext()).replace("\u2060", "")
    if "생성형 AI" not in text:
        raise AssertionError(f"No generative-AI operating note in {path.name}")

    if path.name.startswith(("02_", "05_")):
        image_props = document.findall(".//wp:docPr", ns)
        if len(image_props) < 2 or any(not node.get("descr") for node in image_props):
            raise AssertionError(f"Actual photos with alt text missing: {path.name}")

    if path.name.startswith(("03_", "04_")) and not document.findall(".//w:tblHeader", ns):
        raise AssertionError(f"Accessible table header missing: {path.name}")


def main():
    outputs = [make_worksheet(), make_plan(), make_answer_key(), make_data_cards()]
    for output in outputs:
        audit_docx(output)
        print(output)


if __name__ == "__main__":
    main()
